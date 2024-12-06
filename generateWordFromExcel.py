import pandas as pd
from docx import Document
from docx import Document
from docx.shared import Cm

# Charger le fichier Excel
file_path = 'data/allMetadataCalaU_V2.xlsx'
df = pd.read_excel(file_path)

# Fonction pour nettoyer les en-têtes
def clean_header(header):
    return header.replace('\n', ' ').strip()

# Nettoyer les en-têtes du DataFrame
df.columns = [clean_header(col) for col in df.columns]

# Créer un nouveau document Word
doc = Document()

# Parcourir chaque ligne du DataFrame
for index, row in df.iterrows():
    # Créer un paragraphe avec les informations de la ligne
    paragraphe = f"{row['citation']}\nDurée : {row['duree']}\nDomaine/Mots-clés : {row['domaine_motsCles']}"
    doc.add_paragraph(paragraphe)
    doc.add_paragraph("")  # Ajouter une ligne vide

# Sauvegarder le document Word
output_file = 'Notices_Canal-U.docx'
doc.save(output_file)

print(f"Le fichier Word '{output_file}' a été généré avec succès.")